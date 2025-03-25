import os
import streamlit as st
from tavily import TavilyClient
import google.generativeai as genai
from docx import Document
from io import BytesIO
from config import CONFIG 

# Set Tavily API Key
os.environ["TAVILY_API_KEY"] = CONFIG["TAVILY_API_KEY"]
os.environ["GOOGLE_API_KEY"] = CONFIG["GEMINI_API_KEY"]

# Configure Gemini AI
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])

# Initialize Tavily client
tavily_client = TavilyClient(api_key=os.environ["TAVILY_API_KEY"])

def fetch_industry_insights(industry_name):
    """Fetch latest AI trends & applications in the given industry."""
    query = f"Latest AI applications and trends in {industry_name} industry"
    try:
        response = tavily_client.search(query=query, search_depth="basic", max_results=5)
        results = response.get("results", [])
        
        insights = ""
        for idx, res in enumerate(results, 1):
            insights += f"🔹 **Insight {idx}:** {res['title']}\n"
            insights += f"   - {res['content'][:200]}...\n"  # Show preview
            insights += f"   - 🔗 [Read More]({res['url']})\n\n"
        
        return insights if insights else "No relevant insights found."
    except Exception as e:
        return f"⚠️ Error fetching insights: {str(e)}"

def generate_ai_use_cases_with_gemini(company_name, industry_name):
    """Generates AI Use Case Report using Gemini & Tavily."""
    
    if not company_name or not industry_name:
        return "⚠️ Please provide valid inputs for company and industry."

    # Get industry insights
    industry_insights = fetch_industry_insights(industry_name)

    # Construct prompt for Gemini AI
    prompt = f"""
    You are an AI expert generating AI use cases for companies. Generate structured AI use cases for the company "{company_name}" in the "{industry_name}" industry.
    
    ### Industry Insights:
    {industry_insights}
    
    ### Use Case 1: [Title]
    🔹 **Objective:** Define the problem AI/ML will solve.
    🔹 **AI Application:** Describe how AI/ML will be used.
    🔹 **Business Impact:** Discuss benefits in operations, finance, and supply chain.

    ### Use Case 2: [Title]
    (Repeat the structure for multiple use cases)

    Provide at least 3 use cases.
    """

    # Generate AI response using Gemini
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        use_cases_text = response.text if response and response.text else "No AI use cases generated."
    except Exception as e:
        use_cases_text = f"⚠️ Error generating AI use cases: {str(e)}"

    # Generate DOCX Report in-memory
    doc = Document()
    doc.add_heading(f'GenAI & ML Use Cases for {company_name}', level=1)
    doc.add_paragraph(f"**Industry:** {industry_name}\n\n")

    # Formatting for DOCX
    doc.add_paragraph(f"### Industry Insights\n{industry_insights}\n")
    doc.add_paragraph("### AI Use Cases\n")

    # Splitting Gemini output into bullet points
    use_cases_list = use_cases_text.split("\n")
    for line in use_cases_list:
        if line.strip():
            if line.startswith("### Use Case"):
                doc.add_heading(line.strip(), level=2)
            else:
                doc.add_paragraph(f"🔹 {line.strip()}")

    # Save report to in-memory buffer
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)

    # Display results in a formatted manner
    st.subheader(f"📌 AI Use Cases for {company_name}")
    st.markdown(f"**Industry:** {industry_name}\n\n")
    
    # Display insights with bullet points
    st.markdown("### 🔍 Industry Insights")
    st.markdown(industry_insights)

    # Display use cases in well-structured format
    st.markdown("### 🤖 AI Use Cases")
    st.markdown(use_cases_text.replace("\n", "\n\n"))  # Spacing for readability

    # Provide a download button
    st.download_button("⬇️ Download AI Use Case Report", output_buffer.getvalue(), file_name=f"{company_name.replace(' ', '_')}_AI_Use_Cases.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return use_cases_text
